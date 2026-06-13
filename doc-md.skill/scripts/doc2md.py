#!/usr/bin/env python3
"""Main dispatcher script for document-to-markdown conversion with image extraction."""

import sys
import os
import re
import shutil
import subprocess
from pathlib import Path

SUPPORTED_EXTENSIONS = {'.docx', '.pdf', '.xlsx', '.xls', '.pptx', '.ppt'}

# ── Output Architecture ──────────────────────────────────────────────
# <stem>/              ← named directory = source filename (no ext)
# ├── <stem>.md        ← main Markdown content
# └── images/          ← all extracted images
#     ├── image_1.png
#     └── image_2.png


def prepare_output_dir(input_path: str, output_dir: str = None, output_name: str = None) -> tuple:
    """Create output directory structure: <name>/ and <name>/images/.

    Returns:
        (output_dir, md_path, images_dir)
    """
    input_file = Path(input_path)
    stem = output_name if output_name else input_file.stem

    if output_dir is None:
        output_dir = str(input_file.parent / stem)
    else:
        output_dir = str(Path(output_dir) / stem)

    images_dir = os.path.join(output_dir, "images")
    os.makedirs(images_dir, exist_ok=True)

    md_path = os.path.join(output_dir, f"{stem}.md")
    return output_dir, md_path, images_dir


def check_tool(command):
    """Check if a command-line tool is available."""
    try:
        subprocess.run([command, '--version'], capture_output=True, check=True)
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        return False


def convert_pdf(input_path, output_dir, md_path, images_dir):
    """Convert PDF using PyMuPDF with full structural reconstruction.

    Detects headings (font size), bold/italic (font name + flags),
    code blocks (monospace font), tables (find_tables), lists,
    task lists, and inline image positioning via bbox ordering.
    """
    try:
        import fitz
    except ImportError:
        return None, "PyMuPDF not available. Install: pip install PyMuPDF"

    try:
        result = _convert_pdf_fitz(input_path, md_path, images_dir)
        return result, None
    except Exception as e:
        return None, f"PDF error: {e}"


# ── PDF helper constants ──
BODY_THRESHOLD = 13.0   # font size > this may be a heading
H2_THRESHOLD = 19.0
H1_THRESHOLD = 20.0  # lowered from 23.0 to catch 21.9pt H1
H5_SIZE_TOLERANCE = 1.2  # H5 can be body_size + this tolerance when all-bold

MONO_FONT_MARKERS = ('console', 'courier', 'mono', 'lucida', 'source code', 'monaco', 'dejavu sans mono')
BOLD_MARKERS = ('bold', 'heavy', 'black')
ITALIC_MARKERS = ('italic', 'oblique', 'kursiv')


def _is_bold(font_name, flags):
    return (flags & 1) or any(m in font_name.lower() for m in BOLD_MARKERS)


def _is_italic(font_name, flags):
    return (flags & 2) or any(m in font_name.lower() for m in ITALIC_MARKERS)


def _is_mono(font_name):
    return any(m in font_name.lower() for m in MONO_FONT_MARKERS)


def _heading_level(size, is_bold, body_size=None):
    """Detect heading level from font size. Returns 0 if not a heading."""
    if size >= H1_THRESHOLD and is_bold:
        return 1
    if size >= H2_THRESHOLD:
        return 2
    if size > 16.5:
        return 2
    if size > 14.0:
        return 3
    if size > 12.0:
        return 4
    if size > BODY_THRESHOLD:
        return 5
    # H5 detection: same size as body but all-bold short text
    if body_size and is_bold and size <= body_size + H5_SIZE_TOLERANCE:
        return 5
    return 0


def _span_to_md(span, parent_bold, parent_italic):
    """Convert a text span with formatting to inline markdown."""
    text = span["text"]
    b = _is_bold(span["font"], span["flags"])
    i = _is_italic(span["font"], span["flags"])
    mono = _is_mono(span["font"])

    # Apply formatting wrappers (innermost first)
    if mono:
        text = f"`{text}`"
    if i and not parent_italic:
        text = f"*{text}*"
    if b and not parent_bold:
        text = f"**{text}**"
    return text


def _merge_spans(spans):
    """Merge consecutive same-format spans into groups."""
    if not spans:
        return []
    groups = []
    current = {"spans": [spans[0]], "is_bold": _is_bold(spans[0]["font"], spans[0]["flags"]),
               "is_italic": _is_italic(spans[0]["font"], spans[0]["flags"]),
               "is_mono": _is_mono(spans[0]["font"])}
    for s in spans[1:]:
        b = _is_bold(s["font"], s["flags"])
        i = _is_italic(s["font"], s["flags"])
        m = _is_mono(s["font"])
        if b == current["is_bold"] and i == current["is_italic"] and m == current["is_mono"]:
            current["spans"].append(s)
        else:
            groups.append(current)
            current = {"spans": [s], "is_bold": b, "is_italic": i, "is_mono": m}
    groups.append(current)
    return groups


def _block_text(block):
    """Get full text of a text block."""
    return "".join(s["text"] for line in block["lines"] for s in line["spans"])


def _block_md(block, level, body_size):
    """Convert a text block to markdown, preserving inline formatting."""
    # Collect all spans in order
    span_rows = []
    for line in block["lines"]:
        row_spans = list(line["spans"])
        if row_spans:
            span_rows.append(row_spans)

    if not span_rows:
        return ""

    # Flatten but track line breaks
    flat_spans = []
    for i, row in enumerate(span_rows):
        flat_spans.extend(row)
        if i < len(span_rows) - 1:
            flat_spans.append({"text": "\n", "font": "", "flags": 0})

    # Detect overall block characteristics from first span
    first_size = span_rows[0][0]["size"] if span_rows[0] else body_size
    is_bold_block = all(_is_bold(s["font"], s["flags"]) for row in span_rows for s in row)
    is_italic_block = all(_is_italic(s["font"], s["flags"]) for row in span_rows for s in row)
    all_mono = all(_is_mono(s["font"]) for row in span_rows for s in row)

    # Build inline markdown
    parts = []
    merged = _merge_spans(flat_spans)
    for group in merged:
        group_text = "".join(s["text"] for s in group["spans"])
        if "\n" in group_text and group["is_mono"]:
            parts.append(group_text)
            continue
        if group["is_mono"] and not all_mono:
            group_text = f"`{group_text}`"
        if group["is_italic"] and not is_italic_block:
            group_text = f"*{group_text}*"
        if group["is_bold"] and not is_bold_block:
            group_text = f"**{group_text}**"
        parts.append(group_text)

    text = "".join(parts)

    # Apply block-level formatting
    if level > 0:
        return f"{'#' * level} {text}\n"
    if all_mono:
        return f"```\n{text}\n```\n"
    return f"{text}\n"


def _detect_list(text):
    """Detect list type from text prefix."""
    # Task list
    m = re.match(r'^\s*([\[【])([ xX\s✓✔])[\]】]', text)
    if m:
        checked = 'x' if m.group(2).lower() in ('x', '✓', '✔') else ' '
        return ('task', f"- [{checked}] ")
    # Ordered list
    m = re.match(r'^(\s*)(\d+)[\.\、．)\s]', text)
    if m:
        indent = len(m.group(1))
        num = m.group(2)
        prefix = '  ' * (indent // 2) + '1. '
        return ('ordered', prefix)
    # Unordered list (bullet)
    m = re.match(r'^(\s*)([-•▪▸›»·◦○●◆◇▪▹►])', text)
    if m:
        indent = len(m.group(1))
        prefix = '  ' * max(0, (indent // 2) - 1) + '- '
        if indent < 2:
            prefix = '- '
        return ('unordered', prefix)
    return (None, None)


def _build_list_md(list_items):
    """Convert accumulated list items to markdown with proper nesting."""
    lines = []
    for li_type, prefix, text in list_items:
        # Strip original list marker from text
        if li_type == 'task':
            text = re.sub(r'^\s*[\[【][ xX\s✓✔][\]】]\s*', '', text)
        elif li_type == 'ordered':
            text = re.sub(r'^\s*\d+[\.\、．)\s]\s*', '', text)
        else:
            text = re.sub(r'^\s*[-•▪▸›»·◦○●◆◇▪▹►]\s*', '', text)
        lines.append(f"{prefix}{text}")
    return '\n'.join(lines) + '\n'


def _build_table_md(page, table):
    """Convert a PyMuPDF table to markdown using clip-based text extraction."""
    rows = []
    for row_idx in range(table.row_count):
        cells = []
        for col_idx in range(table.col_count):
            # Get cell text via clip
            cell = table.extract()[row_idx][col_idx] if row_idx < len(table.extract()) else ""
            if cell:
                cell = cell.replace('\n', ' ').strip()
            cells.append(cell)
        rows.append(cells)

    if not rows:
        return ""

    lines = []
    # Header row with separator
    lines.append('| ' + ' | '.join(rows[0]) + ' |')
    lines.append('| ' + ' | '.join(['---'] * len(rows[0])) + ' |')
    for row in rows[1:]:
        lines.append('| ' + ' | '.join(row) + ' |')
    return '\n'.join(lines) + '\n'


def _convert_pdf_fitz(input_path, md_path, images_dir):
    """Core PDF conversion using structured text extraction."""
    import fitz
    doc = fitz.open(input_path)
    global_img_counter = [0]  # mutable counter

    # Phase 1: Detect body font size (most common across all pages)
    all_sizes = []
    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block["type"] == 0:
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        all_sizes.append(span["size"])

    body_size = 9.8  # default
    if all_sizes:
        from collections import Counter
        # Use the most common non-large size
        size_counter = Counter(round(s, 1) for s in all_sizes)
        for s, _ in size_counter.most_common():
            if s < 13:
                body_size = s
                break

    output_parts = []

    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        # Detect tables on this page
        tables = page.find_tables()
        table_bboxes = [t.bbox for t in tables.tables] if hasattr(tables, 'tables') else []
        table_blocks = {i: False for i in range(len(blocks))}  # block inside table?
        for i, block in enumerate(blocks):
            if block["type"] == 0 and "bbox" in block:
                bb = block["bbox"]
                for tb in table_bboxes:
                    if (bb[0] >= tb[0] - 2 and bb[1] >= tb[1] - 2 and
                        bb[2] <= tb[2] + 2 and bb[3] <= tb[3] + 2):
                        table_blocks[i] = True
                        break

        # Extract images with bbox info
        page_images = []
        for block in blocks:
            if block["type"] == 1:
                page_images.append(block["bbox"])

        # Extract actual images from page
        img_xrefs = page.get_images(full=True)
        img_data = []
        for img_info in img_xrefs:
            xref = img_info[0]
            try:
                base_img = doc.extract_image(xref)
                global_img_counter[0] += 1
                ext = base_img['ext']
                fname = f"image_{global_img_counter[0]:03d}.{ext}"
                ipath = os.path.join(images_dir, fname)
                with open(ipath, 'wb') as f:
                    f.write(base_img["image"])
                # Find matching image block bbox
                bbox = page_images.pop(0) if page_images else (0, 0, 0, 0)
                img_data.append((bbox[1], f"![{fname}](images/{fname})"))
            except Exception:
                pass

        # Process text blocks interleaved with images
        img_idx = 0
        page_parts = []
        list_buffer = []  # collect consecutive list items
        code_buffer = []  # collect consecutive code (monospace) blocks

        def flush_list():
            nonlocal list_buffer
            if list_buffer:
                page_parts.append(_build_list_md(list_buffer))
                list_buffer = []

        def flush_code():
            nonlocal code_buffer
            if code_buffer:
                lines = code_buffer
                page_parts.append(f"```\n{chr(10).join(lines)}\n```\n")
                code_buffer = []

        text_blocks = [b for b in blocks if b["type"] == 0 and not table_blocks[blocks.index(b)]]
        # Sort blocks by vertical position
        text_blocks.sort(key=lambda b: b["bbox"][1])

        for block in text_blocks:
            # Insert images that appear above this block
            bbox_top = block["bbox"][1]
            while img_idx < len(img_data) and img_data[img_idx][0] < bbox_top:
                flush_list()
                flush_code()
                page_parts.append(img_data[img_idx][1] + '\n')
                img_idx += 1

            text = _block_text(block)
            if not text.strip():
                # Empty lines flush list/code buffers
                flush_list()
                flush_code()
                page_parts.append('')
                continue

            # Determine block characteristics
            spans = [s for line in block["lines"] for s in line["spans"]]
            if not spans:
                continue

            first_size = spans[0]["size"]
            is_bold_block = all(_is_bold(s["font"], s["flags"]) for s in spans)

            # Count mono spans ratio for soft code detection
            mono_count = sum(1 for s in spans if _is_mono(s["font"]))
            mono_ratio = mono_count / len(spans) if spans else 0
            all_mono = mono_ratio == 1.0

            # Short bold line at body size → potential H5
            text_len = len(text.strip())
            is_short = text_len < 80

            # Heading detection
            level = _heading_level(first_size, is_bold_block, body_size) if is_bold_block or first_size > 13 else 0
            # H5 guard: must be very short isolated text (< 50 chars), not a full sentence
            if level == 5 and (text_len > 50 or first_size > body_size + 1.5):
                level = 0
            # Don't treat long bold paragraphs as headings (level 3+ must be short)
            if level >= 3 and not is_short:
                level = 0

            # Code block detection — soft: majority mono spans
            if mono_ratio >= 0.5 and level == 0:
                flush_list()
                code_buffer.append(text)
                continue

            if mono_ratio < 0.5 and code_buffer:
                flush_code()

            # Heading
            if level > 0:
                flush_list()
                flush_code()
                page_parts.append(f"{'#' * level} {text}\n")
                continue

            # List detection — explicit bullet/number
            li_type, li_prefix = _detect_list(text)
            if li_type:
                flush_code()
                list_buffer.append((li_type, li_prefix, text))
                continue

            # Not a list item, flush list buffer
            if list_buffer:
                flush_list()

            # Regular paragraph with inline formatting
            page_parts.append(_block_md(block, 0, body_size))

        # Flush remaining buffers
        flush_list()
        flush_code()

        # Remaining images
        while img_idx < len(img_data):
            page_parts.append(img_data[img_idx][1] + '\n')
            img_idx += 1

        # Process tables for this page
        for table in tables.tables:
            try:
                table_md = _build_table_md(page, table)
                if table_md:
                    page_parts.append(table_md)
            except Exception:
                pass

        output_parts.append('\n'.join(page_parts))

    doc.close()

    # Final content assembly
    content = '\n'.join(output_parts)
    # Normalize excessive blank lines
    content = re.sub(r'\n{4,}', '\n\n\n', content)
    # Remove trailing whitespace lines
    content = re.sub(r' +\n', '\n', content)

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(content)

    return md_path


def convert_docx(input_path, output_dir, md_path, images_dir):
    """Convert DOCX to Markdown with full rich-text reconstruction.
    
    Strategy:
      - Pandoc (preferred) via subprocess.
      - python-docx fallback with run-level formatting, lists, tables, 
        code blocks, blockquotes, links, and inline image positioning.
    """
    # ── Pandoc path (fast, high quality) ──
    if check_tool('pandoc'):
        try:
            # Run Pandoc from output_dir's parent so paths are relative
            abs_out = os.path.abspath(output_dir)
            abs_md = os.path.abspath(md_path)
            out_name = os.path.basename(output_dir)
            subprocess.run(
                ['pandoc', os.path.abspath(input_path), '-o', out_name + '.md',
                 f'--extract-media={out_name}'],
                cwd=os.path.dirname(abs_out), check=True
            )
            # Pandoc wrote to cwd/<out_name>.md, move to expected path if needed
            pandoc_md = os.path.join(os.path.dirname(abs_out), out_name + '.md')
            if pandoc_md != abs_md and os.path.exists(pandoc_md):
                if os.path.exists(abs_md):
                    os.remove(abs_md)
                shutil.move(pandoc_md, abs_md)
            
            media_dir = os.path.join(output_dir, 'media')
            rId_to_name = {}
            if os.path.isdir(media_dir):
                img_i = 0
                for f in sorted(Path(media_dir).iterdir()):
                    if f.is_file() and f.suffix.lower() in {'.png','.jpg','.jpeg','.gif','.bmp'}:
                        img_i += 1
                        new_name = f"image_{img_i:03d}.{f.suffix.lstrip('.')}"
                        shutil.move(str(f), os.path.join(images_dir, new_name))
                        rId_to_name[f.name] = new_name
                shutil.rmtree(media_dir, ignore_errors=True)
            
            if os.path.exists(md_path):
                raw = Path(md_path).read_text(encoding='utf-8')
                # Fix 1: Strip Pandoc image attributes & rename images
                for old, new in rId_to_name.items():
                    # Replace images: media/rId20.png -> images/image_001.png, strip attrs
                    raw = re.sub(
                        r'!\[([^\]]*)\]\([^)]*?' + re.escape(old) + r'[^)]*\)',
                        r'![\1](images/' + new + r')', raw
                    )
                # Fix 2: Update any remaining media/ path refs
                raw = re.sub(r'!\[([^\]]*)\]\([^)]*?media/([^/)]+)[^)]*\)',
                            r'![\1](images/\2)', raw)
                # Fix 3: Unescape task list brackets: \[x\] -> [x], \[ \] -> [ ]
                raw = re.sub(r'^\s*[-*]\s*\\\[([ xX])\\\]', r'- [\1]', raw, flags=re.MULTILINE)
                Path(md_path).write_text(raw, encoding='utf-8')
            return md_path, None
        except subprocess.CalledProcessError as e:
            return None, f"Pandoc error: {e}"

    # ── python-docx fallback ──
    try:
        from docx import Document
    except ImportError:
        return None, "Neither Pandoc nor python-docx available. Install: pip install python-docx"

    from docx.oxml.ns import qn

    try:
        return _convert_docx_fallback(input_path, md_path, images_dir, Document, qn)
    except Exception as e:
        return None, f"python-docx error: {e}"


def _convert_docx_fallback(input_path, md_path, images_dir, Document, qn):
    """Fallback docx→md using python-docx with full formatting reconstruction."""
    doc = Document(input_path)

    # ── Build numbering info cache ──
    num_info = {}  # numId -> (fmt, level_indent_prefix)
    numbering_part = doc.part.numbering_part
    if numbering_part:
        numbering = numbering_part.element
        abstract_nums = {}
        for abn in numbering.findall(qn('w:abstractNum')):
            an_id = abn.get(qn('w:abstractNumId'))
            levels = {}
            for lvl in abn.findall(qn('w:lvl')):
                ilvl = int(lvl.get(qn('w:ilvl')))
                nf = lvl.find(qn('w:numFmt'))
                fmt = nf.get(qn('w:val')) if nf is not None else 'bullet'
                lt = lvl.find(qn('w:lvlText'))
                lvl_text = lt.get(qn('w:val')) if lt is not None else ''
                levels[ilvl] = {'fmt': fmt, 'lvlText': lvl_text}
            abstract_nums[an_id] = levels
        for num_elem in numbering.findall(qn('w:num')):
            nid = num_elem.get(qn('w:numId'))
            ani = num_elem.find(qn('w:abstractNumId'))
            if ani is not None:
                num_info[nid] = abstract_nums.get(ani.get(qn('w:val')), {})

    # ── Image extraction helper ──
    img_counter = [0]  # mutable counter
    blip_to_ref = {}   # r:embed -> image_N.ext

    def emit_image(embed_id):
        """Save one embedded image to images/ and return markdown ref."""
        if embed_id in blip_to_ref:
            return blip_to_ref[embed_id]
        try:
            rel = doc.part.rels[embed_id]
            if 'image' not in rel.reltype:
                return ''
            img_counter[0] += 1
            image_data = rel.target_part.blob
            ext = rel.target_part.content_type.split('/')[-1]
            if ext == 'jpeg':
                ext = 'jpg'
            fname = f"image_{img_counter[0]:03d}.{ext}"
            with open(os.path.join(images_dir, fname), 'wb') as f:
                f.write(image_data)
            ref = f'![{fname}](images/{fname})'
            blip_to_ref[embed_id] = ref
            return ref
        except Exception:
            return ''

    # ── Paragraph-level image detection ──
    # Map image rId/embed IDs per paragraph so we can insert inline refs
    para_img_refs = {}  # para_index -> markdown ref string
    # Scan all paragraphs for drawing blips
    for pi, para in enumerate(doc.paragraphs):
        for blip in para._element.findall('.//'+qn('a:blip')):
            embed = blip.get(qn('r:embed'))
            if embed:
                ref = emit_image(embed)
                if ref:
                    para_img_refs[pi] = ref

    # ── Run-level text builder with merging ──
    def build_para_text(para):
        """Build markdown text for a paragraph, merging adjacent runs with same formatting."""
        # Check for images in runs first
        for run in para.runs:
            for blip in run._element.findall('.//'+qn('a:blip')):
                embed = blip.get(qn('r:embed'))
                if embed:
                    return emit_image(embed)
        
        # Collect runs, merging adjacent ones with same flags
        merged = []  # list of (text, flags_tuple)
        for run in para.runs:
            t = run.text
            if not t:
                continue
            # Compute formatting flags for this run
            flags = []
            try:
                if run.font.strike: flags.append('S')
            except: pass
            if run.bold: flags.append('B')
            if run.italic: flags.append('I')
            try:
                if run.font.superscript: flags.append('^')
            except: pass
            try:
                if run.font.subscript: flags.append('v')
            except: pass
            flags_key = tuple(flags)
            
            if merged and merged[-1][1] == flags_key:
                # Merge with previous run
                merged[-1] = (merged[-1][0] + t, flags_key)
            else:
                merged.append((t, flags_key))
        
        # Apply markdown formatting to each merged segment
        result = ''
        for text, flags in merged:
            flags_set = set(flags)
            if flags_set == {'B', 'I'}:
                text = f'***{text}***'
            elif 'B' in flags_set:
                text = f'**{text}**'
            elif 'I' in flags_set:
                text = f'*{text}*'
            if 'S' in flags_set:
                text = f'~~{text}~~'
            if '^' in flags_set:
                text = f'<sup>{text}</sup>'
            if 'v' in flags_set:
                text = f'<sub>{text}</sub>'
            result += text
        return result

    # ── List helper ──
    def get_list_marker(para, num_info):
        """Return (prefix, indent_depth, is_ordered) for a list paragraph, or None if not a list."""
        pPr = para._element.find(qn('w:pPr'))
        if pPr is None:
            return None
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return None
        numId_e = numPr.find(qn('w:numId'))
        ilvl_e = numPr.find(qn('w:ilvl'))
        if numId_e is None:
            return None
        nid = numId_e.get(qn('w:val'))
        ilvl = int(ilvl_e.get(qn('w:val'))) if ilvl_e is not None else 0
        info = num_info.get(nid, {})
        level_info = info.get(ilvl, {})
        fmt = level_info.get('fmt', 'bullet')
        indent = '    ' * ilvl
        is_ordered = (fmt == 'decimal')
        return (indent + '-', ilvl, is_ordered)

    # ── Process paragraphs ──
    parts = []
    code_block_lines = []
    blockquote_lines = []
    in_code_block = False
    in_blockquote = False

    def flush_code_block():
        nonlocal code_block_lines, in_code_block
        if code_block_lines and in_code_block:
            # Join all code lines with single newline (not double)
            code_text = '\n'.join(code_block_lines)
            parts.append('```')
            parts.append(code_text)
            parts.append('```')
            parts.append('')
            code_block_lines = []
            in_code_block = False

    def flush_blockquote():
        nonlocal blockquote_lines, in_blockquote
        if blockquote_lines and in_blockquote:
            parts.extend(blockquote_lines)
            parts.append('')
            blockquote_lines = []
            in_blockquote = False

    for pi, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ''

        text = build_para_text(para).strip()

        # ── Heading ──
        heading_match = re.match(r'Heading (\d+)', style_name)
        if heading_match:
            flush_code_block()
            flush_blockquote()
            # If heading has an image, insert image ref after heading
            if pi in para_img_refs:
                parts.append('')
                parts.append(para_img_refs[pi])
                parts.append('')
            if text:
                level = min(int(heading_match.group(1)), 6)
                parts.append(f'{"#" * level} {text}')
            continue

        # ── Source Code ──
        if style_name == 'Source Code':
            flush_blockquote()
            in_code_block = True
            # Normalize line endings: collapse double newlines to single
            code_text = re.sub(r'\n\n+', '\n', para.text.rstrip())
            code_block_lines.append(code_text)
            continue
        else:
            flush_code_block()

        # ── Block Text / Quote ──
        if style_name == 'Block Text':
            flush_code_block()
            in_blockquote = True
            if text:
                blockquote_lines.append(f'> {text}')
            continue
        else:
            flush_blockquote()

        # ── Image-only paragraph (no text, has image ref) ──
        if pi in para_img_refs:
            if text:
                # Text + image on same paragraph — image inline
                parts.append(f'{text}\n\n{para_img_refs[pi]}')
            else:
                parts.append(para_img_refs[pi])
            continue

        # ── Empty paragraph ──
        if not text:
            continue

        # ── List items ──
        list_info = get_list_marker(para, num_info)
        if list_info:
            prefix, ilvl, is_ordered = list_info
            # Check for task list pattern
            m = re.match(r'^\[([ xX])\]\s*(.*)', text)
            if m:
                parts.append(f'{prefix} [{m.group(1)}] {m.group(2)}')
            elif is_ordered:
                # Docx numbering handles the increment — just show text indented
                parts.append(f'{prefix} {text}')
            else:
                parts.append(f'{prefix} {text}')
            continue

        # ── Hyperlinks ──
        hl = para._element.find(qn('w:hyperlink'))
        if hl is not None:
            rid = hl.get(qn('r:id'))
            if rid:
                try:
                    rel = doc.part.rels[rid]
                    url = rel.target_ref
                    hl_runs = []
                    for r in hl.findall('.//'+qn('w:r')):
                        for t_el in r.findall('.//'+qn('w:t')):
                            hl_runs.append(t_el.text or '')
                    link_text = ''.join(hl_runs)
                    if link_text in text:
                        before = text[:text.index(link_text)]
                        after = text[text.index(link_text)+len(link_text):]
                        parts.append(f'{before}[{link_text}]({url}){after}')
                        continue
                except Exception:
                    pass

        # ── Normal paragraph ──
        parts.append(text)

    # Flush any remaining buffered content
    flush_code_block()
    flush_blockquote()

    # ── Tables ──
    if doc.tables:
        parts.append('')
        for table in doc.tables:
            table_md = _table_to_md(table, qn)
            parts.append(table_md)
            parts.append('')

    content = '\n\n'.join(parts)
    # Collapse excessive blank lines
    content = re.sub(r'\n{4,}', '\n\n\n', content)
    Path(md_path).write_text(content, encoding='utf-8')
    return md_path, None


def _fix_code_blocks(raw):
    """Convert Pandoc's 4-space indented code blocks to fenced ``` blocks."""
    lines = raw.split('\n')
    out = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith('    ') and not re.match(r'    [-*+] |    \d+\. ', line):
            if i > 0 and lines[i-1].strip() and not lines[i-1].startswith('```'):
                out.append('```')
                while i < len(lines) and lines[i].startswith('    '):
                    out.append(lines[i][4:])
                    i += 1
                out.append('```')
                out.append('')
                continue
        out.append(line)
        i += 1
    return '\n'.join(out)


def _table_to_md(table, qn):
    """Convert a docx table to Markdown table."""
    # Read grid column count
    tbl_grid = table._tbl.find(qn('w:tblGrid'))
    col_count = len(tbl_grid.findall(qn('w:gridCol'))) if tbl_grid is not None else 1

    rows_data = []
    for row in table.rows:
        cells = []
        col_idx = 0
        for ci, cell in enumerate(row.cells):
            # Get gridSpan
            tc_pr = cell._tc.find(qn('w:tcPr'))
            span = 1
            if tc_pr is not None:
                gs = tc_pr.find(qn('w:gridSpan'))
                if gs is not None:
                    span = int(gs.get(qn('w:val')))
            cell_text = cell.text.replace('\n', ' ').strip()
            cells.append({'text': cell_text, 'span': span, 'col': col_idx})
            col_idx += span
        rows_data.append(cells)

    if not rows_data:
        return ''

    # Build markdown table
    lines = []
    for ri, row in enumerate(rows_data):
        row_texts = []
        for cell in row:
            row_texts.append(cell['text'])
        # Pad to col_count if needed
        while len(row_texts) < col_count:
            row_texts.append('')
        lines.append('| ' + ' | '.join(row_texts) + ' |')
        if ri == 0:
            # Header separator
            sep = '| ' + ' | '.join(['---'] * col_count) + ' |'
            lines.append(sep)

    return '\n'.join(lines)


def convert_excel(input_path, md_path, images_dir):
    """Convert Excel using xlsx2md.py."""
    script_dir = Path(__file__).parent
    script_path = script_dir / 'xlsx2md.py'

    try:
        subprocess.run([sys.executable, str(script_path), input_path, md_path],
                      check=True)
        return md_path, None
    except subprocess.CalledProcessError as e:
        return None, f"Excel conversion error: {e}"


def convert_ppt(input_path, md_path, images_dir):
    """Convert PowerPoint using pptx2md.py."""
    script_dir = Path(__file__).parent
    script_path = script_dir / 'pptx2md.py'

    try:
        subprocess.run([sys.executable, str(script_path), input_path, md_path],
                      check=True)
        return md_path, None
    except subprocess.CalledProcessError as e:
        return None, f"PowerPoint conversion error: {e}"


def main():
    import argparse
    parser = argparse.ArgumentParser(
        description='Convert documents (DOCX/PDF/XLSX/PPTX) to Markdown with image extraction.'
    )
    parser.add_argument('input_file', help='Input document path')
    parser.add_argument('output_dir', nargs='?', default=None, help='Output directory (parent)')
    parser.add_argument('--name', '-n', default=None, help='Custom output name (default: input file stem)')
    args = parser.parse_args()

    input_path = args.input_file
    input_file = Path(input_path)

    if not input_file.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)

    ext = input_file.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        print(f"Error: Unsupported format '{ext}'")
        print(f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}")
        sys.exit(1)

    user_output_dir = args.output_dir
    output_name = args.name

    output_dir, md_path, images_dir = prepare_output_dir(input_path, user_output_dir, output_name)

    # Dispatch based on extension
    if ext == '.docx':
        result, error = convert_docx(input_path, output_dir, md_path, images_dir)
    elif ext == '.pdf':
        result, error = convert_pdf(input_path, output_dir, md_path, images_dir)
    elif ext in ('.xlsx', '.xls'):
        result, error = convert_excel(input_path, md_path, images_dir)
    elif ext in ('.pptx', '.ppt'):
        result, error = convert_ppt(input_path, md_path, images_dir)
    else:
        result, error = None, f"Unhandled format: {ext}"

    if error:
        print(f"Error: {error}")
        sys.exit(1)

    # Summary
    img_count = len([f for f in Path(images_dir).iterdir()]) if Path(images_dir).exists() else 0
    print(f"Converted: {result}")
    print(f"Images: {img_count} ({images_dir})")


if __name__ == "__main__":
    main()
